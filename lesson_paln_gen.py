from docx import Document




# ====== BASIC SETTINGS ======
teacher_name = "Adebisi Victor"
subject = "Animal Husbandry"
class_name = "SS1"
theme = "Farm Animals and Their Body Systems"
term = "First Term"
session = "2025/2026"
output_path = f"{subject.replace(' ', '_')}_{class_name}_{term.replace(' ', '_')}_Lesson_Plans.docx"

# ====== TOPICS (Editable Section) ======
# Format → ("Main Topic", "Sub-Topic")
topics = [
    ("Introduction to Animal Husbandry", ""),
    ("Classification of Farm Animals", ""),
    ("Parts, Organs and Functions in Farm Animals (Digestive System)", ""),
    ("Functions of Parts/Organs of Farm Animals (Circulatory System)", ""),
    ("Functions of Parts/Organs of Farm Animals (Respiratory System)", ""),
    ("Functions of Parts/Organs of Farm Animals (Nervous and Excretory Systems)", ""),
    ("Practical on Organs of Farm Animals (Digestive System)", ""),
    ("Practical on Organs of Farm Animals (Circulatory System)", ""),
    ("Practical on Organs of Farm Animals (Respiratory System)", ""),
    ("Practical on Organs of Farm Animals (Nervous System)", ""),
    ("Practical on Organs of Farm Animals (Excretory System)", "")
]


# ====== LESSON DEVELOPMENT STEPS (Editable Section) ======
# You can reuse this block or change the flow for another subject
steps = [
    ("Introduction (5 mins)", 
     "The teacher asks questions from previous knowledge related to the topic and guides the class to recall relevant experiences.",
     "The students respond based on their prior knowledge or farm experiences and participate in the discussion.",
     "Establishing connection with previous knowledge and interest in the lesson."),

    ("Presentation Step 1 (5 mins)", 
     "The teacher writes out the topic clearly on the board and explains its meaning using simple language.",
     "The students listen attentively, repeat key terms after the teacher, and write the topic in their notebooks.",
     "Understanding the focus and meaning of the topic."),

    ("Presentation Step 2 (7 mins)", 
     "The teacher displays charts, models, or diagrams related to the topic (e.g., organs, body systems, or animal types).",
     "The students observe carefully, identify each item shown, and ask questions for clarification.",
     "Recognition and identification of parts, organs, or animal types."),

    ("Presentation Step 3 (7 mins)", 
     "The teacher explains the structure, function, or classification in detail, providing examples where possible.",
     "The students listen attentively, take notes, and answer questions to show understanding.",
     "Comprehension of the key points and their relationships."),

    ("Presentation Step 4 (7 mins)", 
     "The teacher links the lesson to real-life situations or practical experiences on the school farm.",
     "The students relate what they have learned to real farm practices, sharing their observations or experiences.",
     "Application of knowledge to real-life animal husbandry practices."),

    ("Evaluation (5 mins)", 
     "The teacher asks questions orally or in written form to assess the students’ understanding of the topic.",
     "The students respond confidently to the questions and participate in peer corrections where necessary.",
     "Assessment of the level of understanding achieved."),

    ("Conclusion (4 mins)", 
     "The teacher summarizes the major points of the lesson, corrects any misconceptions, and gives the chalkboard summary.",
     "The students copy the summary neatly into their notebooks and ask final questions for clarification.",
     "Internalization and reinforcement of the lesson content.")
]



# ====== DOCUMENT CREATION FUNCTION ======
def create_lesson_plan():
    doc = Document()
    doc.add_heading(f"SUBJECT: {subject}\tCLASS: {class_name}\tSCHEME OF WORK ({term.upper()})", level=1)

    for i, (topic, subtopic) in enumerate(topics, start=1):
        doc.add_page_break()
        doc.add_heading(f"Week {i} Lesson Plan", level=2)
        doc.add_paragraph(f"Teacher’s Name: {teacher_name}\tTerm: {term}\tSession: {session}")
        doc.add_paragraph(f"\nSubject: {subject}")
        doc.add_paragraph(f"Theme: {theme}")
        doc.add_paragraph(f"Topic: {topic}")
        if subtopic:
            doc.add_paragraph(f"Sub-Topic: {subtopic}")
        doc.add_paragraph("Date:\nTime:\nDuration: 40 minutes\nClass: " + class_name + "\nAverage Age:\nSex: Mixed\nNo. in Class:")

        # Learning Objectives
        doc.add_heading("Learning Objectives:", level=3)
        doc.add_paragraph("At the end of the lesson, the students should be able to:")
        doc.add_paragraph(f"1. Explain the {topic.lower()}.")
        doc.add_paragraph(f"2. Identify key points under {subtopic if subtopic else topic.lower()}.")
        doc.add_paragraph("3. Demonstrate understanding through examples.")

        # Rationale
        doc.add_heading("Rationale / Reason:", level=3)
        doc.add_paragraphrationale_text = (
            "Farm animals play vital roles in human life as sources of food, labour, and raw materials. "
            "It is essential for students to understand the structure and function of different body systems "
            "in farm animals, as this knowledge provides a foundation for effective animal care, management, "
            "and disease prevention."
        )


        # Pre-requisite Knowledge
        doc.add_heading("Pre-requisite Knowledge:", level=3)
        doc.add_paragraph("Students have seen or participated in simple agricultural practices like planting or keeping animals.")

        # Learning Materials
        doc.add_heading("Learning Materials:", level=3)
        doc.add_paragraph("Charts, diagrams, live samples, and farm tools.")

        # Teaching Resources
        doc.add_heading("Teaching Resources:", level=3)
        doc.add_paragraph("School farm, textbooks, and agricultural models.")

        # Reference Material
        doc.add_heading("Reference Material:", level=3)
        doc.add_paragraph("1. Essential Agricultural Science for Senior Secondary Schools.\n"
                          "2. Modern Agricultural Science for Senior Secondary Schools.")

        # Lesson Development
        doc.add_heading("Lesson Development", level=3)
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Stage/Step"
        hdr_cells[1].text = "Teacher’s Activities"
        hdr_cells[2].text = "Students’ Activities"
        hdr_cells[3].text = "Learning Points"

        for step, teacher_act, student_act, learning_point in steps:
            row = table.add_row().cells
            row[0].text = step
            row[1].text = teacher_act
            row[2].text = student_act
            row[3].text = learning_point

    doc.save(output_path)
    print(f"✅ Lesson plan saved as: {output_path}")


# Run it
if __name__ == "__main__":
    create_lesson_plan()
