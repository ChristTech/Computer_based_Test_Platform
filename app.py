import logging
logging.basicConfig(
    filename='app.log',
    level=logging.INFO,
    format='%(asctime)s %(levelname)s: %(message)s'
)

import os
import sqlite3
import uuid
from functools import wraps

import json
import time
from datetime import datetime, timedelta
from flask import Flask, request, render_template, redirect, url_for, jsonify, send_file, abort, session
import hashlib
import random
import docx  # Added for DOCX parsing

import io
import csv
from itsdangerous import URLSafeTimedSerializer, BadSignature, SignatureExpired
import re
from flask import Response

# At the top of your file, after the existing imports
try:
    from docx import Document
    from docx.table import Table
except ImportError:
    Document = None
    Table = None

# Optional dependency for Excel handling
try:
    import pandas as pd
except Exception:
    pd = None

BASE_DIR = os.path.dirname(__file__)
DB = os.path.join(BASE_DIR, 'cbt.db')

# ensure a secret key for session (set a secure value in production)
# store the secret value now and assign it to the Flask app after the app is created
APP_SECRET_KEY = os.environ.get('FLASK_SECRET', 'dev-secret-key-please-change')

# admin creds (env override)
ADMIN_USERNAME = os.environ.get('ADMIN_USERNAME', 'admin')
ADMIN_PASSWORD = os.environ.get('ADMIN_PASSWORD', 'adminpass')

def ensure_column(table: str, column: str, definition: str):
    """
    Ensure a column exists on a table; create it if missing.
    Uses a direct sqlite connection so it can be called before db_conn() is defined.
    """
    try:
        conn = sqlite3.connect(DB)
        cur = conn.cursor()
        cur.execute(f"PRAGMA table_info({table})")
        cols = [row[1] for row in cur.fetchall()]
        if column not in cols:
            cur.execute(f"ALTER TABLE {table} ADD COLUMN {column} {definition}")
            conn.commit()
    except Exception:
        # keep silent; caller may handle errors
        pass
    finally:
        try: conn.close()
        except Exception: pass

def is_admin_request():
    """
    Return True if the request is authenticated as admin.
    Checks session (logged-in admin) first, then accepts admin_password
    provided in JSON body or query string for backward compatibility.
    """
    if 'admin_token' in session:
        return True
    # check JSON payload then query param
    admin_pass = ''
    try:
        if request.json and 'admin_password' in request.json:
            admin_pass = request.json.get('admin_password') or ''
        else:
            admin_pass = request.args.get('admin_password') or ''
    except Exception:
        admin_pass = request.args.get('admin_password') or ''
    return admin_pass == ADMIN_PASSWORD

def init_db():
    conn = sqlite3.connect(DB)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS exams (
        id TEXT PRIMARY KEY,
        title TEXT,
        duration_minutes INTEGER,
        started INTEGER DEFAULT 0,
        teacher_id TEXT
    )''')
    c.execute('''CREATE TABLE IF NOT EXISTS questions (
        id TEXT PRIMARY KEY,
        exam_id TEXT,
        question TEXT,
        choices TEXT,
        answer_index INTEGER
    )''')
    c.execute('''CREATE TABLE IF NOT EXISTS sessions (
        token TEXT PRIMARY KEY,
        exam_id TEXT,
        start_time INTEGER
    )''')
    c.execute('''CREATE TABLE IF NOT EXISTS results (
        id TEXT PRIMARY KEY,
        token TEXT,
        answers TEXT,
        score INTEGER,
        submitted_at INTEGER
    )''')
    c.execute('''CREATE TABLE IF NOT EXISTS registered_students (
        id TEXT PRIMARY KEY,
        exam_id TEXT,
        student_name TEXT
    )''')
    # audit / teacher
    c.execute('''CREATE TABLE IF NOT EXISTS teachers (
        id TEXT PRIMARY KEY,
        name TEXT UNIQUE,
        password_hash TEXT,
        token TEXT UNIQUE
    )''')
    # audit logs
    c.execute('''CREATE TABLE IF NOT EXISTS audit_logs (
        id TEXT PRIMARY KEY,
        ts INTEGER,
        action TEXT,
        teacher_id TEXT,
        exam_id TEXT,
        details TEXT
    )''')
    conn.commit()

    # use module-level helper to ensure optional columns
    ensure_column('results', 'name', 'TEXT')
    ensure_column('results', 'total', 'INTEGER')
    ensure_column('results', 'student_id', 'TEXT')
    ensure_column('sessions', 'student_name', 'TEXT')
    ensure_column('sessions', 'end_time', 'INTEGER')
    ensure_column('exams', 'teacher_id', 'TEXT')
    # audit fields for questions
    ensure_column('questions', 'created_by', 'TEXT')
    ensure_column('questions', 'created_at', 'INTEGER')
    ensure_column('questions', 'updated_by', 'TEXT')
    ensure_column('questions', 'updated_at', 'INTEGER')
    # ensure we can persist per-session ordering
    ensure_column('sessions', 'question_state', 'TEXT')   # JSON: list of {id,question,choices,correct_index}
    ensure_column('sessions', 'question_order', 'TEXT')   # persisted per-session ordering (added)
    ensure_column('results', 'answers_detail', 'TEXT')    # JSON: detailed per-question responses for auditing

    # teacher approval / subject columns (for pending approvals and subject selection)
    ensure_column('teachers', 'approved', 'INTEGER DEFAULT 0')   # 0 = pending, 1 = approved
    ensure_column('teachers', 'subject', 'TEXT')

    # add exam tag and session tag columns (new)
    ensure_column('exams', 'tag', 'TEXT')
    ensure_column('sessions', 'tag', 'TEXT')

    # indexes for faster lookup on larger datasets
    try:
        # c.execute("
        c.execute("CREATE INDEX IF NOT EXISTS idx_regstudents_exam ON registered_students(exam_id)")
        c.execute("CREATE INDEX IF NOT EXISTS idx_questions_exam ON questions(exam_id)")
        conn.commit()
    except Exception:
        pass

    conn.close()

app = Flask(__name__)
# assign the previously read secret to the Flask app now that `app` exists
app.secret_key = APP_SECRET_KEY
init_db()

def db_conn():
    conn = sqlite3.connect(DB)
    conn.row_factory = sqlite3.Row
    return conn

def _hash_password(password: str) -> str:
    return hashlib.sha256((password or '').encode('utf-8')).hexdigest()

def get_teacher_by_token(token):
    if not token: return None
    conn = db_conn(); c = conn.cursor()
    c.execute('SELECT id,name,token FROM teachers WHERE token=?', (token,))
    t = c.fetchone()
    conn.close()
    return t

def get_teacher_from_request():
    token = request.headers.get('X-Teacher-Token') or (request.json and request.json.get('teacher_token'))
    if not token: return None
    return get_teacher_by_token(token)

@app.route('/')
def index():
    return redirect(url_for('login_page'))

@app.route('/admin')
def admin():
    if 'admin_token' not in session:
        return redirect(url_for('login_page'))
    return render_template('admin.html')

@app.route('/teacher')
def teacher():
    return render_template('teacher.html')

@app.route('/student')
def student_landing():
    return render_template('student.html')

@app.route('/login')
def login_page():
    return render_template('login.html')

@app.route('/api/create_exam', methods=['POST'])
def create_exam():
    data = request.json or {}
    title = data.get('title') or 'Exam'
    try:
        duration = int(data.get('duration') or 30)
    except Exception:
        duration = 30
    tag = (data.get('tag') or '').strip() or None
    exam_id = str(uuid.uuid4())[:8]
    conn = db_conn(); c = conn.cursor()
    c.execute('INSERT INTO exams (id,title,duration_minutes,tag) VALUES (?,?,?,?)', (exam_id, title, duration, tag))
    conn.commit(); conn.close()
    return jsonify({'exam_id': exam_id})

@app.route('/api/add_question', methods=['POST'])
def add_question():
    data = request.json or {}
    teacher = get_teacher_from_request()
    if not teacher:
        return jsonify({'error': 'teacher auth required'}), 401
    exam_id = data.get('exam_id')
    q = data.get('question')
    choices = data.get('choices') or []
    try:
        answer_index = int(data.get('answer_index') or 0)
    except Exception:
        answer_index = 0
    if not exam_id or not q:
        return jsonify({'error': 'missing fields'}), 400
    conn = db_conn(); c = conn.cursor()
    c.execute('SELECT teacher_id FROM exams WHERE id=?', (exam_id,))
    er = c.fetchone()
    if not er:
        conn.close(); return jsonify({'error': 'exam not found'}), 400
    if er['teacher_id'] and er['teacher_id'] != teacher['id']:
        conn.close(); return jsonify({'error': 'not allowed to add questions for this exam'}), 403
    qid = str(uuid.uuid4())[:8]
    now = int(time.time())
    c.execute('INSERT INTO questions (id,exam_id,question,choices,answer_index,created_by,created_at) VALUES (?,?,?,?,?,?,?)',
              (qid, exam_id, q, json.dumps(choices), answer_index, teacher['id'], now))
    conn.commit(); conn.close()
    log_audit('add_question', teacher['id'], exam_id, {'question_id': qid, 'answer_index': answer_index})
    return jsonify({'ok': True, 'question_id': qid})

@app.route('/api/list_exams')
def list_exams():
    conn = db_conn(); c = conn.cursor()
    # include teacher.subject and exam.tag for UI convenience
    c.execute('''
        SELECT e.id, e.title, e.duration_minutes, e.started, e.teacher_id,
               COALESCE(t.subject, '') AS subject,
               COALESCE(e.tag, '') AS tag
        FROM exams e
        LEFT JOIN teachers t ON e.teacher_id = t.id
        ORDER BY subject DESC
    ''')
    rows = c.fetchall(); conn.close()
    out = []
    for r in rows:
        out.append({
            'id': r['id'],
            'title': r['title'],
            'duration_minutes': r['duration_minutes'],
            'started': bool(r['started']),
            'teacher_id': r['teacher_id'],
            'subject': r['subject'] or '',
            'tag': r['tag'] or ''
        })
    return jsonify(out)

@app.route('/api/questions/<exam_id>')
def get_questions(exam_id):
    conn = db_conn(); c = conn.cursor()
    c.execute('SELECT id,question,choices,answer_index FROM questions WHERE exam_id=?', (exam_id,))
    rows = c.fetchall(); conn.close()
    qs = []
    for r in rows:
        qs.append({'id': r['id'], 'question': r['question'], 'choices': json.loads(r['choices'] or '[]'), 'answer_index': r['answer_index']})
    return jsonify(qs)

@app.route('/api/upload_questions', methods=['POST'])
def upload_questions():
    """
    Accept file (CSV, Excel, or DOCX) and exam_id via form-data.
    If form field 'overwrite' == '1', existing questions for the exam are deleted first.
    Teachers may upload questions for their own exams by sending X-Teacher-Token.
    """
    exam_id = request.form.get('exam_id')
    file = request.files.get('file')
    teacher = get_teacher_from_request()
    overwrite = request.form.get('overwrite') == '1'

    if not exam_id or not file:
        return jsonify({'error': 'Missing exam_id or file'}), 400

    # quick validation (avoid large processing if exam missing / permission denied)
    conn = db_conn(); c = conn.cursor()
    c.execute('SELECT id, teacher_id FROM exams WHERE id=?', (exam_id,))
    er = c.fetchone()
    if not er:
        conn.close(); return jsonify({'error': 'exam not found'}), 400
    if er['teacher_id'] and (not teacher or teacher['id'] != er['teacher_id']):
        conn.close(); return jsonify({'error': 'not allowed to upload questions for this exam'}, 403)

    # parse file content to entries: (question, choices[], answer_index)
    entries = []
    fname = (file.filename or '').lower()
    ext = os.path.splitext(fname)[1]

    try:
        if ext in ('.xlsx', '.xls') and pd:
            df = pd.read_excel(file)
            df.columns = [str(col).strip().lower() for col in df.columns]
            for _, row in df.iterrows():
                question = row.get('question')
                if not question or (isinstance(question, float) and pd.isna(question)):
                    continue
                choices = []
                for i in range(1, 10):
                    for key in (f'choice{i}', f'option{i}'):
                        val = row.get(key)
                        if val and not (isinstance(val, float) and pd.isna(val)):
                            choices.append(str(val))
                if not choices:
                    # try columns named choice*
                    for col in df.columns:
                        if str(col).lower().startswith('choice') and pd.notna(row.get(col)):
                            choices.append(str(row.get(col)))
                answer_index = 0
                if 'answer_index' in row.index and pd.notna(row.get('answer_index')):
                    try: answer_index = int(row.get('answer_index'))
                    except Exception: answer_index = 0
                entries.append((str(question).strip(), choices, int(answer_index)))
        elif ext == '.docx' and Document:
            doc = Document(file)
            current_question = None
            choices = []
            answer_index = None
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    if current_question and len(choices) >= 2 and answer_index is not None:
                        entries.append((current_question, choices, answer_index))
                        current_question = None; choices = []; answer_index = None
                    continue
                if text.lower().startswith('question:'):
                    if current_question and len(choices) >= 2 and answer_index is not None:
                        entries.append((current_question, choices, answer_index))
                    current_question = text[len('question:'):].strip()
                    choices = []; answer_index = None
                elif any(text.lower().startswith(p) for p in ('a.', 'b.', 'c.', 'd.')):
                    # "A. choice"
                    choice_text = text.split('.', 1)[1].strip() if '.' in text else text
                    choices.append(choice_text)
                elif text.lower().startswith('answer:'):
                    a = text[len('answer:'):].strip().upper()
                    amap = {'A':0,'B':1,'C':2,'D':3}
                    answer_index = amap.get(a, 0)
            if current_question and len(choices) >= 2 and answer_index is not None:
                entries.append((current_question, choices, answer_index))
            # fallback: try tables
            if not entries:
                for table in doc.tables:
                    headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
                    if 'question' not in headers: continue
                    for row in table.rows[1:]:
                        row_data = [cell.text.strip() for cell in row.cells]
                        try:
                            question = row_data[headers.index('question')]
                        except Exception:
                            continue
                        choices = []
                        for i, h in enumerate(headers):
                            if h.startswith(('choice', 'option')) and row_data[i]:
                                choices.append(row_data[i])
                        answer_index = 0
                        if 'answer_index' in headers:
                            ai = row_data[headers.index('answer_index')]
                            try: answer_index = int(ai)
                            except Exception:
                                try:
                                    amap = {'a':0,'b':1,'c':2,'d':3}
                                    answer_index = amap.get(str(ai).strip().lower(), 0)
                                except Exception:
                                    answer_index = 0
                        if question and len(choices) >= 2:
                            entries.append((question, choices, answer_index))
        else:
            # CSV / text
            stream = io.StringIO(file.stream.read().decode('utf-8', errors='ignore'))
            reader = csv.DictReader(stream)
            for r in reader:
                question = r.get('question') or None
                if not question:
                    continue
                choices = []
                for k, v in r.items():
                    if k and str(k).lower().startswith(('choice', 'option')) and v:
                        choices.append(v)
                if not choices and r.get('choices'):
                    choices = [c.strip() for c in r.get('choices').split('|') if c.strip()] or [c.strip() for c in r.get('choices').split(',') if c.strip()]
                try:
                    answer_index = int(r.get('answer_index') or 0)
                except Exception:
                    ans = (r.get('answer') or '').strip().upper()
                    amap = {'A':0,'B':1,'C':2,'D':3}
                    answer_index = amap.get(ans, 0)
                entries.append((question, choices, answer_index))
    except Exception as e:
        conn.close()
        app.logger.exception("upload parsing error: %s", e)
        return jsonify({'error': f'Failed parsing file: {str(e)}'}), 400

    if not entries:
        conn.close()
        return jsonify({'error': 'No valid questions found in file. Ensure correct format.'}), 400

    c = conn.cursor()
    inserted = 0
    if overwrite:
        c.execute('DELETE FROM questions WHERE exam_id=?', (exam_id,))
    for question, choices, answer_index in entries:
        if not question or len(choices) < 2:
            continue
        qid = str(uuid.uuid4())[:8]
        now = int(time.time())
        c.execute('INSERT INTO questions (id,exam_id,question,choices,answer_index,created_by,created_at) VALUES (?,?,?,?,?,?,?)',
                  (qid, exam_id, question, json.dumps(choices), int(answer_index), teacher['id'] if teacher else None, now))
        inserted += 1
        log_audit('add_question', teacher['id'] if teacher else None, exam_id, {'question_id': qid, 'answer_index': answer_index})
    conn.commit(); conn.close()
    app.logger.info("Uploaded %d questions to exam %s by teacher=%s", inserted, exam_id, teacher['id'] if teacher else '-')

    # single audit event for the bulk upload action (notifies admin via audit_logs)
    try:
        log_audit('upload_questions', teacher['id'] if teacher else None, exam_id, {'count': inserted, 'overwrite': bool(overwrite)})
    except Exception:
        pass

    return jsonify({'ok': True, 'count': inserted})

@app.route('/api/start_exam', methods=['POST'])
def start_exam():
    data = request.json or {}
    exam_id = data.get('exam_id')
    student_name = data.get('student_name') or data.get('name')
    if not exam_id:
        return jsonify({'error': 'exam_id required'}), 400

    conn = db_conn(); c = conn.cursor()
    c.execute('SELECT id, duration_minutes, started, tag FROM exams WHERE id=?', (exam_id,))
    ex = c.fetchone()
    if not ex:
        conn.close(); return jsonify({'error': 'exam not found'}), 404
    if not ex['started']:
        conn.close(); return jsonify({'error': 'exam not open'}), 403

    # fetch questions and build per-session presentation state
    c.execute('SELECT id, question, choices, answer_index FROM questions WHERE exam_id=?', (exam_id,))
    qrows = c.fetchall()
    if not qrows:
        conn.close(); return jsonify({'error': 'no questions for exam'}), 400

    qlist = []
    for r in qrows:
        choices = json.loads(r['choices'] or '[]')
        qlist.append({'id': r['id'], 'question': r['question'], 'choices': choices, 'answer_index': int(r['answer_index'] or 0)})

    # randomize question order
    random.shuffle(qlist)

    # for each question, shuffle choices and compute presented correct_index
    question_state = []
    for q in qlist:
        choices = list(q['choices'])
        random.shuffle(choices)
        # find presented correct index by matching value of original correct choice
        orig_correct_idx = int(q.get('answer_index', 0))
        orig_choice_value = (q['choices'][orig_correct_idx] if orig_correct_idx < len(q['choices']) else None)
        presented_correct = 0
        if orig_choice_value is not None:
            for i, val in enumerate(choices):
                if val == orig_choice_value:
                    presented_correct = i
                    break
        question_state.append({
            'id': q['id'],
            'question': q['question'],
            'choices': choices,
            'correct_index': presented_correct
        })

    start_time = int(time.time())
    duration = int(ex['duration_minutes'] or 30) * 60
    end_time = start_time + duration
    token = str(uuid.uuid4())[:8]
    exam_tag = ex['tag'] if 'tag' in ex.keys() else None

    c.execute('INSERT INTO sessions (token, exam_id, start_time, end_time, student_name, question_state, tag) VALUES (?,?,?,?,?,?,?)',
              (token, exam_id, start_time, end_time, student_name or None, json.dumps(question_state), exam_tag))
    conn.commit(); conn.close()

    url = request.host_url.rstrip('/') + url_for('exam_page', token=token)
    return jsonify({'token': token, 'url': url})
# (moved ensure_column('sessions', 'question_order', 'TEXT') into init_db to avoid calling an undefined local function)
ensure_column('sessions', 'question_order', 'TEXT')

@app.route('/exam/<token>')
def exam_page(token):
    conn = db_conn(); c = conn.cursor()
    c.execute('SELECT exam_id,start_time,end_time,question_state FROM sessions WHERE token=?', (token,))
    row = c.fetchone()
    if not row:
        conn.close(); abort(404, description="Invalid token")
    exam_id = row['exam_id']
    end_time = row['end_time'] if 'end_time' in row.keys() else None
    qs = []
    try:
        qs = json.loads(row['question_state'] or '[]')
    except Exception:
        qs = []
    conn.close()
    remaining = max(0, int(end_time - int(time.time()))) if end_time else 0
    # render page with questions exactly as stored
    return render_template('exam.html', token=token, title=exam_id, questions=qs, remaining_seconds=remaining)

@app.route('/api/submit/<token>', methods=['POST'])
def submit(token):
    data = request.json or {}
    answers = data.get('answers') or {}
    name = (data.get('name') or '').strip()
    conn = db_conn(); c = conn.cursor()
    c.execute('SELECT exam_id, start_time, end_time, student_name FROM sessions WHERE token=?', (token,))
    row = c.fetchone()
    if not row:
        conn.close(); return jsonify({'error': 'invalid token'}), 400
    # enforce time limit
    now = int(time.time())
    if 'end_time' in row.keys() and row['end_time'] and now > row['end_time']:
        conn.close(); return jsonify({'error': 'exam time expired'}), 403
    exam_id = row['exam_id']
    if not name and 'student_name' in row.keys() and row['student_name']:
        name = row['student_name']

    # load the per-session question_state (presentation order + shuffled choices)
    c.execute('SELECT question_state FROM sessions WHERE token=?', (token,))
    srow = c.fetchone()
    try:
        qstate = json.loads(srow['question_state'] or '[]') if srow else []
    except Exception:
        qstate = []

    # build detailed answers and compute score using session's correct_index
    score = 0
    answers_detail = []
    for q in qstate:
        qid = q.get('id')
        qtext = q.get('question') or ''
        choices = q.get('choices') or []
        correct_index = int(q.get('correct_index') or 0)

        # accept either numeric index or letter (a/b/c...) from client
        sel_idx = None
        try:
            raw = answers.get(qid)
            if raw is None or raw == '':
                sel_idx = None
            elif isinstance(raw, (int, float)):
                sel_idx = int(raw)
            else:
                # could be "A" / "a" / "B" / "0" etc.
                s = str(raw).strip()
                if s.isdigit():
                    sel_idx = int(s)
                elif len(s) >= 1 and s[0].isalpha():
                    sel_idx = ord(s[0].upper()) - 65
                else:
                    sel_idx = None
        except Exception:
            sel_idx = None

        sel_text = (choices[sel_idx] if (sel_idx is not None and 0 <= sel_idx < len(choices)) else '')
        correct_text = (choices[correct_index] if 0 <= correct_index < len(choices) else '')
        is_correct = (sel_idx is not None and sel_idx == correct_index)
        if is_correct:
            score += 1

        # add letter labels for easier cross-checking (A/B/C/...)
        sel_label = (chr(65 + sel_idx) if (sel_idx is not None and 0 <= sel_idx < 26) else None)
        correct_label = (chr(65 + correct_index) if 0 <= correct_index < 26 else None)

        answers_detail.append({
            'id': qid,
            'question': qtext,
            'choices': choices,
            'selected_index': sel_idx,
            'selected_label': sel_label,
            'selected_text': sel_text,
            'correct_index': correct_index,
            'correct_label': correct_label,
            'correct_text': correct_text,
            'is_correct': bool(is_correct)
        })

    submitted_at = int(time.time())
    rid = str(uuid.uuid4())[:8]
    c.execute('DELETE FROM results WHERE token=?', (token,))
    try:
        c.execute('INSERT INTO results (id,token,name,answers,answers_detail,score,submitted_at,total) VALUES (?,?,?,?,?,?,?,?)',
                  (rid, token, name, json.dumps(answers), json.dumps(answers_detail), score, submitted_at, len(qstate)))
    except sqlite3.OperationalError:
        c.execute('INSERT INTO results (id,token,answers,score,submitted_at) VALUES (?,?,?,?,?)',
                  (rid, token, json.dumps(answers), score, submitted_at))
    conn.commit(); conn.close()
    try:
        save_result_to_excel(name, token, exam_id, score, len(qstate), submitted_at, answers_detail)
    except Exception:
        pass

    # audit: student submission event for admin review/notifications
    try:
        # try to include subject for this exam in audit details
        subj = ''
        try:
            conn = db_conn(); c = conn.cursor()
            c.execute('SELECT t.subject FROM exams e LEFT JOIN teachers t ON e.teacher_id=t.id WHERE e.id=?', (exam_id,))
            er = c.fetchone()
            subj = (er['subject'] or '').strip() if er and 'subject' in er.keys() else ''
            conn.close()
        except Exception:
            subj = ''
        log_audit('submit_exam', None, exam_id, {'token': token, 'name': name or '', 'score': score, 'total': len(qstate), 'subject': subj})
    except Exception:
        pass

    return jsonify({'score': score, 'total': len(qstate)})

def _sanitize_filename(s: str) -> str:
    s = (s or '').strip()
    if not s:
        return 'unknown'
    s = s.lower()
    s = re.sub(r'\s+', '_', s)
    s = re.sub(r'[^a-z0-9_\-\.]', '', s)
    return s[:120] or 'unknown'

def save_result_to_excel(name, token, exam_id, score, total, submitted_at, answers_detail=None):
    # determine subject label (try DB -> exams -> teacher.subject), fallback to exam_id
    subject_name = ''
    exam_tag = ''
    try:
        conn = db_conn(); c = conn.cursor()
        c.execute('SELECT t.subject, e.tag FROM exams e LEFT JOIN teachers t ON e.teacher_id=t.id WHERE e.id=?', (exam_id,))
        er = c.fetchone()
        subject_name = (er['subject'] or '').strip() if er and 'subject' in er.keys() else ''
        exam_tag = (er['tag'] or '').strip() if er and 'tag' in er.keys() else ''
        conn.close()
    except Exception:
        subject_name = ''
        exam_tag = ''

    subject_label = _sanitize_filename(subject_name or exam_id)
    tag_label = _sanitize_filename(exam_tag or '')

    row = {
        'tag': exam_tag or '',
        'subject': subject_name or '',
        'name': name or '',
        'token': token,
        'score': score,
        'total': total,
        'submitted_at': datetime.fromtimestamp(submitted_at).isoformat()
    }

    fname_xlsx = os.path.join(BASE_DIR, f'results_{subject_label}.xlsx')
    fname_csv = os.path.join(BASE_DIR, f'results_{subject_label}.csv')
    # also maintain per-tag files for quick downloads
    fname_tag_xlsx = os.path.join(BASE_DIR, f'results_tag_{tag_label}.xlsx') if tag_label else None
    fname_tag_csv = os.path.join(BASE_DIR, f'results_tag_{tag_label}.csv') if tag_label else None

    # include answers_detail if present
    if answers_detail is not None:
        row['answers_detail'] = json.dumps(answers_detail, ensure_ascii=False)

    # Try XLSX (pandas) and persist to disk so admin can download by subject file
    if pd:
        try:
            df_row = pd.DataFrame([row])
            if os.path.exists(fname_xlsx):
                try:
                    existing = pd.read_excel(fname_xlsx)
                    df = pd.concat([existing, df_row], ignore_index=True)
                except Exception:
                    df = df_row
            else:
                df = df_row
            df.to_excel(fname_xlsx, index=False)
            if tag_label:
                try:
                    if os.path.exists(fname_tag_xlsx):
                        existing_tag = pd.read_excel(fname_tag_xlsx)
                        df_tag = pd.concat([existing_tag, df_row], ignore_index=True)
                    else:
                        df_tag = df_row
                    df_tag.to_excel(fname_tag_xlsx, index=False)
                except Exception:
                    pass
            # also ensure CSV exists for systems that expect CSV
            try:
                df.to_csv(fname_csv, index=False)
                if tag_label:
                    try: df.to_csv(fname_tag_csv, index=False)
                    except Exception: pass
            except Exception:
                pass
            return
        except Exception:
            app.logger.exception("failed to write xlsx for subject %s", subject_label)

    # CSV fallback / append
    write_header = not os.path.exists(fname_csv)
    fieldnames = ['subject','name','token','score','total','submitted_at']
    if tag_label:
        fieldnames = ['tag'] + fieldnames
    if answers_detail is not None:
        fieldnames.append('answers_detail')
    with open(fname_csv, 'a', newline='', encoding='utf-8') as fh:
        writer = csv.DictWriter(fh, fieldnames=fieldnames)
        if write_header:
            writer.writeheader()
        writer.writerow(row)
    # append to tag CSV as well
    if tag_label:
        write_header_t = not os.path.exists(fname_tag_csv)
        with open(fname_tag_csv, 'a', newline='', encoding='utf-8') as fh2:
            writer2 = csv.DictWriter(fh2, fieldnames=fieldnames)
            if write_header_t:
                writer2.writeheader()
            writer2.writerow(row)

@app.route('/api/results_csv/<exam_id>')
def results_csv(exam_id):
    # endpoint disabled temporarily while we rework export logic
    return jsonify({'error': 'results_csv disabled temporarily'}), 410

@app.route('/api/upload_students', methods=['POST'])
def upload_students():
    exam_id = request.form.get('exam_id')
    file = request.files.get('file')
    if not exam_id or not file:
        return jsonify({'error': 'Missing exam_id or file'}), 400
    conn = db_conn(); c = conn.cursor()
    c.execute('SELECT id FROM exams WHERE id=?', (exam_id,))
    if not c.fetchone():
        conn.close(); return jsonify({'error': 'exam not found'}), 400
    names = []
    fname = file.filename.lower()
    try:
        if pd and (fname.endswith('.xlsx') or fname.endswith('.xls')):
            df = pd.read_excel(file)
            for _, row in df.iterrows():
                n = row.get('name') or row.get('student') or None
                if n and not (isinstance(n, float) and pd.isna(n)):
                    names.append(str(n).strip())
        else:
            stream = io.StringIO(file.stream.read().decode('utf-8'))
            reader = csv.DictReader(stream)
            for r in reader:
                n = r.get('name') or next(iter(r.values()), None)
                if n:
                    names.append(n.strip())
    except Exception:
        conn.close(); return jsonify({'error': 'Failed parsing file'}), 400

    inserted = 0
    for name in names:
        if not name: continue
        c.execute('SELECT id FROM registered_students WHERE exam_id=? AND LOWER(student_name)=LOWER(?)', (exam_id, name))
        if c.fetchone(): continue
        sid = str(uuid.uuid4())[:8]
        c.execute('INSERT INTO registered_students (id,exam_id,student_name) VALUES (?,?,?)', (sid, exam_id, name))
        inserted += 1
    conn.commit(); conn.close()
    return jsonify({'ok': True, 'count': inserted})

@app.route('/api/add_student', methods=['POST'])
def add_student():
    data = request.json or {}
    exam_id = (data.get('exam_id') or '').strip()
    name = (data.get('name') or '').strip()
    if not exam_id or not name:
        return jsonify({'error': 'exam_id and name required'}), 400
    conn = db_conn(); c = conn.cursor()
    c.execute('SELECT id FROM exams WHERE id=?', (exam_id,))
    if not c.fetchone():
        conn.close(); return jsonify({'error': 'exam not found'}), 400
    c.execute('SELECT id FROM registered_students WHERE exam_id=? AND LOWER(student_name)=LOWER(?)', (exam_id, name))
    if c.fetchone():
        conn.close(); return jsonify({'ok': True, 'note': 'already registered'})
    sid = str(uuid.uuid4())[:8]
    c.execute('INSERT INTO registered_students (id,exam_id,student_name) VALUES (?,?,?)', (sid, exam_id, name))
    conn.commit(); conn.close()
    return jsonify({'ok': True, 'student_id': sid})

@app.route('/api/list_students/<exam_id>')
def list_students(exam_id):
    conn = db_conn(); c = conn.cursor()
    c.execute('SELECT id,student_name FROM registered_students WHERE exam_id=? ORDER BY student_name', (exam_id,))
    rows = c.fetchall(); conn.close()
    return jsonify([{'id': r['id'], 'name': r['student_name']} for r in rows])

@app.route('/results/<token>')
def results_page(token):
    conn = db_conn(); c = conn.cursor()
    c.execute('SELECT r.score, r.submitted_at, r.name, s.exam_id FROM results r JOIN sessions s ON r.token=s.token WHERE r.token=?', (token,))
    row = c.fetchone(); conn.close()
    if not row:
        return "Result not found", 404
    score, submitted_at, name, exam_id = row['score'], row['submitted_at'], row['name'] or '', row['exam_id']
    return render_template('results.html', score=score, submitted_at=submitted_at, exam_id=exam_id, name=name, countdown=180)

@app.route('/api/register_teacher', methods=['POST'])
def register_teacher():
    data = request.json or {}
    name = (data.get('name') or '').strip()
    password = data.get('password') or ''
    subject = (data.get('subject') or '').strip()
    if not name or not password:
        return jsonify({'error': 'name and password required'}), 400
    conn = db_conn(); c = conn.cursor()
    c.execute('SELECT id FROM teachers WHERE name=?', (name,))
    if c.fetchone():
        conn.close(); return jsonify({'error': 'teacher exists'}), 400
    tid = str(uuid.uuid4())[:8]
    ph = _hash_password(password)
    # create teacher as PENDING (approved=0). Admin will approve later.
    c.execute('INSERT INTO teachers (id,name,password_hash,approved,subject,token) VALUES (?,?,?,?,?,?)', (tid, name, ph, 0, subject or None, None))
    conn.commit(); conn.close()

    # audit: teacher registration (include subject)
    try:
        log_audit('teacher_registered', tid, None, {'name': name, 'subject': subject})
    except Exception:
        pass

    return jsonify({'ok': True, 'note': 'pending_approval', 'name': name})

@app.route('/api/pending_teachers')
def pending_teachers():
    conn = db_conn(); c = conn.cursor()
    # return teachers awaiting admin approval
    try:
        c.execute('SELECT id, name, subject FROM teachers WHERE approved=0 ORDER BY subject DESC')
        rows = c.fetchall()
    finally:
        conn.close()
    return jsonify([{'id': r['id'], 'name': r['name'], 'subject': r['subject']} for r in rows])

@app.route('/api/teacher_student_scores', methods=['GET'])
def teacher_student_scores():
    teacher = get_teacher_from_request()
    if not teacher:
        return jsonify({'error': 'Teacher authentication required'}), 401

    conn = db_conn()
    c = conn.cursor()

    # fetch results for exams belonging to this teacher
    c.execute('''
        SELECT r.name, r.score, r.total, r.submitted_at, e.title as exam_title, e.id as exam_id
        FROM results r
        JOIN sessions s ON r.token = s.token
        JOIN exams e ON s.exam_id = e.id
        WHERE e.teacher_id = ?
        ORDER BY e.id, r.submitted_at DESC
    ''', (teacher['id'],))
    rows = c.fetchall()

    # collect question counts per exam to compute percentages when total missing
    exam_ids = sorted({row['exam_id'] for row in rows})
    q_counts = {}
    for eid in exam_ids:
        c.execute('SELECT COUNT(1) as cnt FROM questions WHERE exam_id=?', (eid,))
        qr = c.fetchone()
        q_counts[eid] = qr['cnt'] if qr and 'cnt' in qr.keys() else 0

    # format rows
    scores_json = []
    for row in rows:
        total = row['total'] if ('total' in row.keys() and row['total'] is not None) else q_counts.get(row['exam_id'], 0)
        submitted_at = None
        try:
            submitted_at = datetime.fromtimestamp(row['submitted_at']).strftime('%Y-%m-%d %H:%M') if row['submitted_at'] else ''
        except Exception:
            submitted_at = ''
        score_val = row['score'] if row['score'] is not None else 0
        percentage = round((score_val / total) * 100, 1) if total and total > 0 else 0
        scores_json.append({
            'name': row['name'] or 'Anonymous',
            'score': score_val,
            'total': total,
            'percentage': percentage,
            'submitted_at': submitted_at,
            'exam_title': row['exam_title'],
            'exam_id': row['exam_id']
        })

    conn.close()

    # group by exam for frontend convenience
    exams = {}
    for s in scores_json:
        aid = s['exam_id']
        if aid not in exams:
            exams[aid] = {'title': s['exam_title'], 'scores': []}
        exams[aid]['scores'].append(s)

    # compute overall stats
    overall_avg = round(sum(s['score'] for s in scores_json) / len(scores_json), 1) if scores_json else 0

    return jsonify({
        'ok': True,
        'teacher': teacher['name'],
        'exams': list(exams.values()),
        'total_students': len(scores_json),
        'overall_avg': overall_avg
    })

@app.route('/api/login_teacher', methods=['POST'])
def login_teacher():
    data = request.json or {}
    name = (data.get('name') or '').strip()
    password = data.get('password') or ''
    if not name or not password:
        return jsonify({'error': 'name and password required'}), 400
    ph = _hash_password(password)
    conn = db_conn(); c = conn.cursor()
    c.execute('SELECT id,token FROM teachers WHERE name=? AND password_hash=?', (name, ph))
    row = c.fetchone()
    conn.close()
    if not row:
        return jsonify({'error': 'invalid credentials'}), 401
    return jsonify({'ok': True, 'teacher_token': row['token'], 'name': name})

@app.route('/api/teacher_create_exam', methods=['POST'])
def teacher_create_exam():
    data = request.json or {}
    teacher = get_teacher_from_request()
    if not teacher:
        return jsonify({'error': 'teacher auth required'}), 401
    title = (data.get('title') or 'Exam').strip()
    try:
        duration = int(data.get('duration') or 30)
    except Exception:
        duration = 30
    tag = (data.get('tag') or '').strip() or None
    exam_id = str(uuid.uuid4())[:8]
    conn = db_conn(); c = conn.cursor()
    c.execute('INSERT INTO exams (id,title,duration_minutes,teacher_id,tag) VALUES (?,?,?,?,?)', (exam_id, title, duration, teacher['id'], tag))
    conn.commit(); conn.close()
    return jsonify({'ok': True, 'exam_id': exam_id})

@app.route('/api/set_exam_state', methods=['POST'])
def set_exam_state():
    data = request.json or {}
    exam_id = data.get('exam_id')
    started = bool(data.get('started'))
    if not is_admin_request():
        return jsonify({'error': 'admin auth required'}), 401
    if not exam_id:
        return jsonify({'error': 'exam_id required'}), 400
    conn = db_conn(); c = conn.cursor()
    c.execute('UPDATE exams SET started=? WHERE id=?', (1 if started else 0, exam_id))
    conn.commit(); conn.close()
    return jsonify({'ok': True, 'exam_id': exam_id, 'started': started})

def log_audit(action, teacher_id, exam_id, details=None):
    try:
        conn = db_conn(); c = conn.cursor()
        aid = str(uuid.uuid4())[:8]
        c.execute('INSERT INTO audit_logs (id,ts,action,teacher_id,exam_id,details) VALUES (?,?,?,?,?,?)',
                  (aid, int(time.time()), action, teacher_id, exam_id, json.dumps(details or {})))
        conn.commit(); conn.close()
    except Exception:
        pass

@app.after_request
def set_security_headers(response):
    # basic security headers
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'DENY'
    response.headers['Referrer-Policy'] = 'no-referrer-when-downgrade'
    response.headers['X-XSS-Protection'] = '1; mode=block'
    # minimal CSP - adjust as needed
    response.headers['Content-Security-Policy'] = "default-src 'self' 'unsafe-inline' data:;"
    return response

# Admin decorator must be defined before any route uses it
def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'admin_token' not in session:
            return jsonify({'error': 'Admin authentication required'}), 401
        return f(*args, **kwargs)
    return decorated_function

# Admin login / session endpoints (restore if missing)
@app.route('/api/login_admin', methods=['POST'])
def login_admin():
    """
    Accepts JSON or form: { username, password }.
    Sets session['admin_token'] on success.
    """
    data = {}
    try:
        data = request.get_json(silent=True) or {}
    except Exception:
        data = {}
    # also accept form-encoded fallback
    if not data:
        data = { 'username': request.form.get('username'), 'password': request.form.get('password') }
    username = (data.get('username') or '').strip()
    password = data.get('password') or ''
    if username == ADMIN_USERNAME and password == ADMIN_PASSWORD:
        session['admin_token'] = uuid.uuid4().hex
        return jsonify({'ok': True})
    return jsonify({'error': 'Invalid credentials'}), 401

@app.route('/api/check_admin')
def check_admin():
    return jsonify({'is_admin': 'admin_token' in session})

@app.route('/api/logout_admin', methods=['POST','GET'])
def logout_admin():
    session.pop('admin_token', None)
    return jsonify({'ok': True})

# Add a user-facing logout endpoint (templates often call url_for('logout'))
@app.route('/logout')
def logout():
    """
    Clear admin session and redirect to the login page.
    Provides the 'logout' endpoint expected by templates.
    """
    session.pop('admin_token', None)
    return redirect(url_for('login_page'))

@app.errorhandler(404)
def not_found(e):
    return render_template('base_error.html', code=404, message=str(e)), 404

@app.errorhandler(500)
def server_error(e):
    return render_template('base_error.html', code=500, message='Internal server error'), 500

# audit logs JSON endpoint
@app.route('/api/audit_logs')
def api_audit_logs():
    """
    Return audit logs, grouped by date. Supports date filter.
    If format=csv or format=xlsx and date=YYYY-MM-DD is provided,
    return a downloadable file containing rows: name, score for submit_exam actions.
    """
    exam_id = request.args.get('exam_id')
    teacher_id = request.args.get('teacher_id')
    date_str = request.args.get('date')  # YYYY-MM-DD
    fmt = (request.args.get('format') or '').lower()

    conn = db_conn(); c = conn.cursor()
    q = "SELECT ts, action, teacher_id, exam_id, details FROM audit_logs"
    params = []
    where_clauses = []

    if exam_id:
        where_clauses.append("exam_id=?"); params.append(exam_id)
    if teacher_id:
        where_clauses.append("teacher_id=?"); params.append(teacher_id)
    if date_str:
        try:
            date_obj = datetime.strptime(date_str, '%Y-%m-%d').date()
            start_ts = int(datetime(date_obj.year, date_obj.month, date_obj.day).timestamp())
            end_ts = int(datetime(date_obj.year, date_obj.month, date_obj.day, 23, 59, 59).timestamp())
            where_clauses.append("ts BETWEEN ? AND ?")
            params.extend([start_ts, end_ts])
        except ValueError:
            conn.close()
            return jsonify({'error': 'Invalid date format. Use YYYY-MM-DD'}), 400

    if where_clauses:
        q += " WHERE " + " AND ".join(where_clauses)
    q += " ORDER BY ts DESC"
    c.execute(q, tuple(params))
    rows = c.fetchall(); conn.close()

    # If client requested a file for a specific date, produce name+score file
    if fmt in ('csv', 'xlsx') and date_str:
        records = []
        for r in rows:
            try:
                details = json.loads(r['details'] or '{}')
            except Exception:
                details = {}
            if r['action'] == 'submit_exam':
                name = details.get('name') or ''
                score = details.get('score')
                # only include entries that have numeric score
                try:
                    score_val = int(score) if score is not None else ''
                except Exception:
                    score_val = score
                records.append({'name': name, 'score': score_val})

        # build download
        if fmt == 'xlsx' and pd:
            try:
                df = pd.DataFrame(records)
                mem = io.BytesIO()
                with pd.ExcelWriter(mem, engine='openpyxl') as writer:
                    df.to_excel(writer, index=False, sheet_name='audit_logs')
                mem.seek(0)
                return send_file(mem,
                                 mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                                 as_attachment=True,
                                 download_name=f'audit_logs_{date_str}.xlsx')
            except Exception:
                app.logger.exception("failed to build xlsx audit logs for %s", date_str)
                # fallthrough to CSV
        # CSV fallback
        si = io.StringIO()
        writer = csv.writer(si)
        writer.writerow(['name', 'score'])
        for rec in records:
            writer.writerow([rec.get('name',''), rec.get('score','')])
        mem = io.BytesIO(); mem.write(si.getvalue().encode('utf-8')); mem.seek(0)
        return send_file(mem, mimetype='text/csv', as_attachment=True, download_name=f'audit_logs_{date_str}.csv')

    # Default behavior: grouped JSON by date
    grouped_logs = {}
    for r in rows:
        try:
            details = json.loads(r['details'] or '{}')
        except Exception:
            details = {}
        log_entry = {'ts': r['ts'], 'action': r['action'], 'teacher_id': r['teacher_id'], 'exam_id': r['exam_id'], 'details': details}
        date = datetime.fromtimestamp(r['ts']).strftime('%Y-%m-%d')
        grouped_logs.setdefault(date, []).append(log_entry)

    return jsonify(grouped_logs)

@app.route('/admin/logs')
def admin_logs():
    return render_template('admin_logs.html')

@app.route('/api/results/<token>')
def api_get_result(token):
    conn = db_conn(); c = conn.cursor()
    c.execute('SELECT r.name, r.score, r.total, r.answers, r.submitted_at, s.exam_id FROM results r JOIN sessions s ON r.token=s.token WHERE r.token=?', (token,))
    row = c.fetchone()
    conn.close()
    if not row:
        return jsonify({'error': 'result not found'}), 404
    answers = []
    try:
        answers = json.loads(row['answers'] or '[]')
    except Exception:
        answers = []
    # prepare minimal answer breakdown if possible
    resp = {
        'name': row['name'] or '',
        'score': row['score'] if row['score'] is not None else 0,
        'total': row['total'] if 'total' in row.keys() and row['total'] is not None else (len(answers) or 0),
        'answers': answers,
        'submitted_at': row['submitted_at'],
        'exam_id': row['exam_id'],
        'can_retake': False
    }
    return jsonify(resp)

@app.route('/api/results_all')
def results_all():
    # endpoint disabled temporarily while we rework export logic
    return jsonify({'error': 'results export endpoint disabled temporarily'}), 410

@app.route('/api/subjects')
def api_subjects():
    """
    Return a JSON list of subjects loaded from subjects.xlsx or subjects.csv in project root.
    """
    subjects = []
    csv_path = os.path.join(BASE_DIR, 'subjects.csv')
    xlsx_path = os.path.join(BASE_DIR, 'subjects.xlsx')
    try:
        # prefer Excel if available and pandas installed
        if pd and os.path.exists(xlsx_path):
            df = pd.read_excel(xlsx_path)
            cols = [str(c).strip().lower() for c in df.columns]
            if 'subject' in cols:
                col = df.columns[cols.index('subject')]
                subjects = [str(x).strip() for x in df[col].dropna().astype(str).tolist()]
            else:
                subjects = [str(x).strip() for x in df.iloc[:, 0].dropna().astype(str).tolist()]
        elif os.path.exists(csv_path):
            with open(csv_path, encoding='utf-8') as f:
                rdr = csv.reader(f)
                header = next(rdr, None)
                if header:
                    header_l = [h.strip().lower() for h in header]
                    if 'subject' in header_l:
                        idx = header_l.index('subject')
                        for row in rdr:
                            if len(row) > idx and row[idx].strip():
                                subjects.append(row[idx].strip())
                    else:
                        # treat first column as subject name (include header first-col if non-empty)
                        if header and header[0].strip():
                            subjects.append(header[0].strip())
                        for row in rdr:
                            if row and row[0].strip():
                                subjects.append(row[0].strip())
                else:
                    for row in rdr:
                        if row and row[0].strip():
                            subjects.append(row[0].strip())
    except Exception:
        app.logger.exception("failed to load subjects")
    return jsonify(subjects)

@app.route('/api/subjects_db')
def api_subjects_db():
    """
    Return distinct subjects from the database (teachers/exams). If none found,
    fall back to /api/subjects (file-based).
    """
    subjects = []
    try:
        conn = db_conn(); c = conn.cursor()
        # prefer teacher.subject
        c.execute("SELECT DISTINCT TRIM(subject) AS subject FROM teachers WHERE subject IS NOT NULL AND TRIM(subject) <> '' ORDER BY subject")
        rows = c.fetchall()
        subjects = [r['subject'] for r in rows if r['subject']]
        # if empty, try subjects from exams' teachers
        if not subjects:
            c.execute('''
                SELECT DISTINCT TRIM(t.subject) AS subject
                FROM exams e
                JOIN teachers t ON e.teacher_id = t.id
                WHERE t.subject IS NOT NULL AND TRIM(t.subject) <> ''
                ORDER BY subject
            ''')
            rows = c.fetchall()
            subjects = [r['subject'] for r in rows if r['subject']]
        conn.close()
    except Exception:
        app.logger.exception("api_subjects_db failed")
        subjects = []

    # fallback to file-based endpoint if DB has none
    if not subjects:
        try:
            return api_subjects()
        except Exception:
            pass

    return jsonify(subjects)

@app.route('/api/approve_teacher', methods=['POST'])
def approve_teacher():
    """
    Admin approves a pending teacher.
    Body JSON: { "teacher_id": "..." }
    Accepts session-based admin or admin_password in request for compatibility.
    Returns teacher_token on success.
    """
    data = request.json or {}
    teacher_id = data.get('teacher_id')
    if not is_admin_request():
        return jsonify({'error': 'admin auth required'}), 401
    if not teacher_id:
        return jsonify({'error': 'teacher_id required'}), 400

    conn = db_conn(); c = conn.cursor()
    c.execute('SELECT id, approved FROM teachers WHERE id=?', (teacher_id,))
    row = c.fetchone()
    if not row:
        conn.close(); return jsonify({'error': 'teacher not found'}), 404
    if row['approved']:
        conn.close(); return jsonify({'error': 'already approved'}), 400

    token = str(uuid.uuid4())[:24]
    try:
        c.execute('UPDATE teachers SET approved=1, token=? WHERE id=?', (token, teacher_id))
        conn.commit()
    finally:
        conn.close()

    try:
        log_audit('teacher_approved', teacher_id, None, {'token': token})
    except Exception:
        pass

    return jsonify({'ok': True, 'teacher_token': token})

@app.route('/api/results_json/<exam_id>')
def results_json(exam_id):
    # endpoint disabled temporarily while we rework export/view logic
    return jsonify({'error': 'results JSON endpoint disabled temporarily'}), 410


@app.route('/api/download_exam_results/<exam_id>')
def download_exam_results(exam_id):
    # endpoint disabled temporarily while we rework export logic
    return jsonify({'error': 'download_exam_results disabled temporarily'}), 410

@app.route('/api/download_subject')
@admin_required
def download_subject():
    """
    Download all results for a given subject.
    Accepts: ?subject=<name>&exam_id=<id>&format=xlsx|csv
    If exam_id is provided we use it to resolve the subject (more reliable).
    """
    subject = (request.args.get('subject') or '').strip()
    exam_id = (request.args.get('exam_id') or '').strip()
    fmt = (request.args.get('format') or 'csv').lower()

    # If exam_id provided, try to resolve subject from the exam -> teacher
    if exam_id:
        try:
            conn = db_conn(); c = conn.cursor()
            c.execute('SELECT t.subject, e.tag FROM exams e LEFT JOIN teachers t ON e.teacher_id=t.id WHERE e.id=?', (exam_id,))
            er = c.fetchone()
            if er:
                # prefer teacher.subject if present
                subj_from_exam = (er['subject'] or '').strip()
                if subj_from_exam:
                    subject = subj_from_exam
            conn.close()
        except Exception:
            pass

    if not subject:
        return jsonify({'error': 'subject required (or exam_id must resolve a subject)'}), 400

    label = _sanitize_filename(subject)
    fname_xlsx = os.path.join(BASE_DIR, f'results_{label}.xlsx')
    fname_csv = os.path.join(BASE_DIR, f'results_{label}.csv')

    # prefer an existing per-subject file on disk
    if fmt == 'xlsx' and pd and os.path.exists(fname_xlsx):
        return send_file(fname_xlsx,
                         mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                         as_attachment=True,
                         download_name=f'{label}.xlsx')

    if os.path.exists(fname_csv):
        return send_file(fname_csv,
                         mimetype='text/csv',
                         as_attachment=True,
                         download_name=f'{label}.csv')

    # Build from DB (same as before)
    conn = db_conn(); c = conn.cursor()
    q = '''
        SELECT r.token, r.name, r.score, r.total, r.submitted_at, r.answers_detail, e.id AS exam_id, COALESCE(t.subject,'') AS subject, COALESCE(e.tag,'') AS tag
        FROM results r
        JOIN sessions s ON r.token = s.token
        LEFT JOIN exams e ON s.exam_id = e.id
        LEFT JOIN teachers t ON e.teacher_id = t.id
        WHERE LOWER(TRIM(COALESCE(t.subject,''))) = ?
        ORDER BY r.submitted_at DESC
    '''
    try:
        c.execute(q, (subject.lower(),))
        rows = c.fetchall()
    finally:
        conn.close()

    # prepare csv/xlsx stream and return (unchanged)...
    data = []
    for r in rows:
        try:
            answers_detail = json.loads(r['answers_detail'] or '[]')
        except Exception:
            answers_detail = r['answers_detail'] or ''
        submitted = ''
        try:
            submitted = datetime.fromtimestamp(r['submitted_at']).isoformat() if r['submitted_at'] else ''
        except Exception:
            submitted = ''
        data.append({
            'exam_id': r['exam_id'] or '',
            'token': r['token'],
            'name': r['name'] or '',
            'score': r['score'],
            'total': r['total'] if 'total' in r.keys() else '',
            'submitted_at': submitted,
            'answers_detail': json.dumps(answers_detail, ensure_ascii=False)
        })

    # XLSX if requested and pandas available
    if fmt == 'xlsx' and pd:
        try:
            df = pd.DataFrame(data)
            mem = io.BytesIO()
            with pd.ExcelWriter(mem, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='results')
            mem.seek(0)
            return send_file(mem,
                             mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                             as_attachment=True,
                             download_name=f'{label}.xlsx')
        except Exception:
            app.logger.exception("failed to build xlsx for %s", label)

    # CSV fallback
    si = io.StringIO()
    writer = csv.writer(si)
    writer.writerow(['exam_id','token','name','score','total','submitted_at','answers_detail'])
    for d in data:
        writer.writerow([d['exam_id'], d['token'], d['name'], d['score'], d['total'], d['submitted_at'], d['answers_detail']])
    mem = io.BytesIO(); mem.write(si.getvalue().encode('utf-8')); mem.seek(0)
    return send_file(mem, mimetype='text/csv', as_attachment=True, download_name=f'{label}.csv')

# serializer for signed temporary downloads
# uses app.secret_key; ensure app.secret_key is set earlier in file
def _get_download_serializer():
    secret = getattr(app, 'secret_key', None) or os.environ.get('FLASK_SECRET', 'dev-secret')
    return URLSafeTimedSerializer(secret, salt='cbt-download-salt')

# Create a short-lived signed download link (admin-only)
@app.route('/api/create_download_link')
@admin_required
def create_download_link():
    subject = (request.args.get('subject') or '').strip()
    tag = (request.args.get('tag') or '').strip()
    fmt = (request.args.get('format') or 'csv').lower()
    if not subject and not tag:
        return jsonify({'error': 'subject or tag required'}), 400
    payload = {'subject': subject, 'tag': tag, 'format': fmt}
    s = _get_download_serializer()
    token = s.dumps(payload)
    url = url_for('download_token', token=token, _external=True)
    return jsonify({'ok': True, 'url': url})

# Public download endpoint that accepts a signed token (valid for a short time)
@app.route('/download/<token>')
def download_token(token):
    s = _get_download_serializer()
    max_age = int(os.environ.get('DOWNLOAD_TOKEN_TTL', '300'))  # seconds, default 5min
    try:
        payload = s.loads(token, max_age=max_age)
    except SignatureExpired:
        return jsonify({'error': 'download link expired'}), 410
    except BadSignature:
        return jsonify({'error': 'invalid download link'}), 400

    subject = payload.get('subject') or ''
    tag = payload.get('tag') or ''
    fmt = (payload.get('format') or 'csv').lower()

    # resolve label (prefer subject, else tag)
    label_key = subject or tag
    if not label_key:
        return jsonify({'error': 'no subject/tag in token'}), 400
    label = _sanitize_filename(label_key)

    # prefer saved files on disk
    fname_xlsx = os.path.join(BASE_DIR, f'results_{label}.xlsx')
    fname_csv = os.path.join(BASE_DIR, f'results_{label}.csv')
    # tag-specific files use prefix results_tag_
    fname_tag_xlsx = os.path.join(BASE_DIR, f'results_tag_{label}.xlsx')
    fname_tag_csv = os.path.join(BASE_DIR, f'results_tag_{label}.csv')

    if fmt == 'xlsx' and pd:
        for path in (fname_xlsx, fname_tag_xlsx):
            if path and os.path.exists(path):
                return send_file(path,
                                 mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                                 as_attachment=True,
                                 download_name=os.path.basename(path))
    # csv fallback / if requested csv
    for path in (fname_csv, fname_tag_csv):
        if path and os.path.exists(path):
            return send_file(path, mimetype='text/csv', as_attachment=True, download_name=os.path.basename(path))

    # build from DB: choose by subject first, else by tag
    conn = db_conn(); c = conn.cursor()
    try:
        if subject:
            q = '''
            SELECT r.token, r.name, r.score, r.total, r.submitted_at, r.answers_detail, e.id AS exam_id, COALESCE(t.subject,'') AS subject, COALESCE(e.tag,'') AS tag
            FROM results r
            JOIN sessions s ON r.token = s.token
            LEFT JOIN exams e ON s.exam_id = e.id
            LEFT JOIN teachers t ON e.teacher_id = t.id
            WHERE LOWER(TRIM(COALESCE(t.subject,''))) = ?
            ORDER BY r.submitted_at DESC
            '''
            c.execute(q, (subject.lower(),))
        else:
            q = '''
            SELECT r.token, r.name, r.score, r.total, r.submitted_at, r.answers_detail, e.id AS exam_id, COALESCE(t.subject,'') AS subject, COALESCE(e.tag,'') AS tag
            FROM results r
            JOIN sessions s ON r.token = s.token
            LEFT JOIN exams e ON s.exam_id = e.id
            LEFT JOIN teachers t ON e.teacher_id = t.id
            WHERE LOWER(TRIM(COALESCE(e.tag,''))) = ?
            ORDER BY r.submitted_at DESC
            '''
            c.execute(q, (tag.lower(),))
        rows = c.fetchall()
    finally:
        conn.close()

    data = []
    for r in rows:
        try:
            answers_detail = json.loads(r['answers_detail'] or '[]')
        except Exception:
            answers_detail = r['answers_detail'] or ''
        submitted = ''
        try:
            submitted = datetime.fromtimestamp(r['submitted_at']).isoformat() if r['submitted_at'] else ''
        except Exception:
            submitted = ''
        data.append({
            'exam_id': r['exam_id'] or '',
            'token': r['token'],
            'name': r['name'] or '',
            'score': r['score'],
            'total': r['total'] if 'total' in r.keys() else '',
            'submitted_at': submitted,
            'answers_detail': json.dumps(answers_detail, ensure_ascii=False)
        })

    # produce xlsx if requested and pandas present
    if fmt == 'xlsx' and pd:
        try:
            df = pd.DataFrame(data)
            mem = io.BytesIO()
            with pd.ExcelWriter(mem, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='results')
            mem.seek(0)
            return send_file(mem,
                             mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                             as_attachment=True,
                             download_name=f'{label}.xlsx')
        except Exception:
            app.logger.exception("failed to build xlsx for %s", label)

    # csv fallback
    si = io.StringIO()
    writer = csv.writer(si)
    writer.writerow(['exam_id','token','name','score','total','submitted_at','answers_detail'])
    for d in data:
        writer.writerow([d['exam_id'], d['token'], d['name'], d['score'], d['total'], d['submitted_at'], d['answers_detail']])
    mem = io.BytesIO(); mem.write(si.getvalue().encode('utf-8')); mem.seek(0)
    return send_file(mem, mimetype='text/csv', as_attachment=True, download_name=f'{label}.csv')

@app.template_filter('datetime')
def format_datetime(value):
    if not value:
        return ''
    if isinstance(value, (int, float)):
        dt = datetime.fromtimestamp(value)
        return dt.strftime('%Y-%m-%d %H:%M:%S')
    return str(value)

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('FLASK_DEBUG', '1') == '1'
    # bind to all interfaces so remote devices (or Docker) can reach it; change host if needed
    app.run(host='0.0.0.0', port=port, debug=debug)